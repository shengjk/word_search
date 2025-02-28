import os
import time
import psutil
import docx
from pdfminer.high_level import extract_text
from collections import defaultdict
import jieba
from cache_manager import CacheManager
import logger_config

logger = logger_config.setup_logger(__name__)

def process_document(file_path, doc_type, timeout=180, cache_manager=None):
    try:
        logger.info(f"\n[文档处理] 开始处理{doc_type}文档")
        logger.info(f"[文档处理] 文件路径: {file_path}")
        start_time = time.time()

        # 检查缓存
        if cache_manager:
            cached_doc = cache_manager.get_cached_document(file_path)
            if cached_doc:
                logger.info(f"[文档处理] 找到缓存数据，直接使用缓存")
                return cached_doc
        text = ""

        # 检查文件大小并设置动态超时
        file_size = os.path.getsize(file_path) / (1024 * 1024)  # 转换为MB
        logger.info(f"[文档处理] 文件大小: {file_size:.2f}MB")
        if file_size > 100:  # 如果文件大于100MB
            logger.info(f"[文档处理] 警告: 文件大小超过100MB，可能需要较长处理时间")
            timeout = min(timeout * 2, 600)  # 大文件增加超时时间，但不超过10分钟

        # 获取当前内存使用情况
        process = psutil.Process()
        memory_info = process.memory_info()
        initial_memory = memory_info.rss / 1024 / 1024
        logger.info(f"[文档处理] 初始内存使用: {initial_memory:.2f}MB")

        # 添加内存监控函数
        def check_memory_usage():
            current_memory = process.memory_info().rss / 1024 / 1024
            memory_increase = current_memory - initial_memory
            if memory_increase > 1024:  # 如果内存增加超过1GB
                raise MemoryError(f"内存使用增加过多: {memory_increase:.2f}MB")
            if time.time() - start_time > timeout:
                raise TimeoutError(f"处理超时（{timeout}秒）")

        if doc_type == 'docx':
            logger.info("\n[Word文档处理] 开始读取文件...")
            logger.info(f"[Word文档处理] 文件路径: {file_path}")
            # 使用内存映射优化Word文档读取
            with open(file_path, 'rb') as f:
                logger.info("[Word文档处理] 正在加载文档对象...")
                doc = docx.Document(f)
                text_parts = []
                total_paragraphs = len(doc.paragraphs)
                logger.info(f"[Word文档处理] 文档加载完成，总段落数: {total_paragraphs}")
                
                for i, paragraph in enumerate(doc.paragraphs):
                    if time.time() - start_time > timeout:
                        raise TimeoutError(f"处理超时（{timeout}秒）")
                    if paragraph.text.strip():  # 只添加非空段落
                        text_parts.append(paragraph.text)
                    if i % 50 == 0:  # 每处理50个段落输出一次进度
                        progress = (i / total_paragraphs) * 100
                        elapsed_time = time.time() - start_time
                        memory_info = process.memory_info()
                        logger.info(f"[Word文档处理] 进度: {progress:.1f}% ({i}/{total_paragraphs} 段落)")
                        logger.info(f"[Word文档处理] 已用时间: {elapsed_time:.1f}秒")
                        logger.info(f"[Word文档处理] 当前内存使用: {memory_info.rss / 1024 / 1024:.1f}MB")
                        if i > 0:
                            avg_time_per_para = elapsed_time / i
                            remaining_paras = total_paragraphs - i
                            estimated_remaining_time = avg_time_per_para * remaining_paras
                            logger.info(f"[Word文档处理] 预计剩余时间: {estimated_remaining_time:.1f}秒")
                
                text = '\n'.join(text_parts)
                logger.info(f"\n[Word文档处理] 文档读取完成")
                logger.info(f"[Word文档处理] 有效段落数: {len(text_parts)}")
                logger.info(f"[Word文档处理] 文本总长度: {len(text)} 字符")
                del doc, text_parts  # 释放内存
        else:  # pdf
            logger.info("\n[PDF文档处理] 开始读取文件...")
            logger.info(f"[PDF文档处理] 文件路径: {file_path}")
            # 使用pdfminer.six优化PDF处理
            # 检查文件大小，对大文件采取特殊处理
            if file_size > 50:  # 如果文件大于50MB
                logger.info(f"[PDF文档处理] 大文件处理模式")
                timeout = min(timeout, 300)  # 限制最大超时时间为5分钟

            logger.info("[PDF文档处理] 正在提取文本...")
            extract_start_time = time.time()
            try:
                with open(file_path, 'rb') as pdf_file:
                    text = extract_text(
                        pdf_file,
                        maxpages=50,  # 限制最大页数
                        caching=True,  # 启用缓存
                        codec='utf-8'
                    )
                    if not text:
                        raise ValueError("无法提取文本内容")
                    
                extract_time = time.time() - extract_start_time
                logger.info(f"[PDF文档处理] 文本提取耗时: {extract_time:.1f}秒")
                logger.info(f"[PDF文档处理] 文本提取完成，文本长度: {len(text)} 字符")
            except Exception as e:
                logger.info(f"[PDF文档处理] 文本提取失败: {str(e)}")
                return None

        logger.info("\n[分词处理] 开始进行分词...")
        text_lower = text.lower()
        logger.info(f"[分词处理] 文本预处理完成，准备分词")
        batch_size = min(5000, max(1000, int(1000000 / file_size)))  # 根据文件大小动态调整批处理大小
        words = []
        word_positions = defaultdict(list)
        
        # 使用生成器优化内存使用
        def word_generator(text):
            logger.info("[分词处理] 初始化分词生成器...")
            for i, word in enumerate(jieba.cut(text)):
                if i % 1000 == 0:  # 增加检查频率
                    check_memory_usage()
                yield i, word

        # 分批处理分词
        word_gen = word_generator(text_lower)
        current_batch = []
        total_words = 0
        batch_count = 0

        logger.info("[分词处理] 开始批量处理...")
        segment_start_time = time.time()
        while True:
            try:
                if time.time() - start_time > timeout:
                    raise TimeoutError(f"处理超时（{timeout}秒）")
                
                i, word = next(word_gen)
                current_batch.append((i, word))
                
                if len(current_batch) >= batch_size:
                    batch_count += 1
                    batch_start_time = time.time()
                    
                    for idx, w in current_batch:
                        words.append(w)
                        word_positions[w].append(idx)
                    total_words += len(current_batch)
                    
                    batch_time = time.time() - batch_start_time
                    segment_time = time.time() - segment_start_time
                    memory_info = process.memory_info()
                    
                    logger.info(f"\n[分词处理] 批次 {batch_count} 处理完成:")
                    logger.info(f"[分词处理] - 本批处理耗时: {batch_time:.2f}秒")
                    logger.info(f"[分词处理] - 总耗时: {segment_time:.2f}秒")
                    logger.info(f"[分词处理] - 当前已处理: {total_words} 个词")
                    logger.info(f"[分词处理] - 当前词典大小: {len(word_positions)} 个不同词")
                    logger.info(f"[分词处理] - 当前内存使用: {memory_info.rss / 1024 / 1024:.1f}MB")
                    logger.info(f"[分词处理] - 平均处理速度: {total_words/segment_time:.1f} 词/秒")
                    
                    current_batch = []
                    
            except StopIteration:
                # 处理最后一批
                if current_batch:
                    for idx, w in current_batch:
                        words.append(w)
                        word_positions[w].append(idx)
                    total_words += len(current_batch)
                    logger.info("\n[分词处理] 处理最后一批数据完成")
                break

        process_time = time.time() - start_time
        memory_info = process.memory_info()
        
        logger.info(f"\n[处理完成] 文档处理结果汇总:")
        logger.info(f"[处理完成] - 总用时: {process_time:.2f}秒")
        logger.info(f"[处理完成] - 总词数: {len(words)}")
        logger.info(f"[处理完成] - 不同词数: {len(word_positions)}")
        logger.info(f"[处理完成] - 最终内存使用: {memory_info.rss / 1024 / 1024:.1f}MB")
        logger.info(f"[处理完成] - 平均处理速度: {len(words)/process_time:.1f} 词/秒")

        result = {
            'path': str(file_path),
            'content': text,
            'type': doc_type,
            'word_positions': dict(word_positions),
            'words': words,
            'process_time': process_time
        }

        # 保存到缓存
        if cache_manager:
            cache_manager.cache_document(file_path, result)
            logger.info(f"[文档处理] 文档处理结果已保存到缓存")

        return result
    except Exception as e:
        logger.info(f"\n错误: 处理{doc_type}文档 {file_path} 失败")
        logger.info(f"错误信息: {str(e)}")
        logger.info(f"错误类型: {type(e).__name__}")
        return None