import sys
from PyQt6.QtWidgets import QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, \
    QPushButton, QLineEdit, QLabel, QFileDialog, QTextEdit, QProgressBar
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QTextCharFormat, QColor, QTextCursor
from PyQt6.QtWidgets import (QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
                             QLineEdit, QPushButton, QProgressBar, QTextEdit,
                             QFileDialog)
from PyQt6.QtGui import QTextCursor, QTextCharFormat, QColor, QTextDocument
from PyQt6.QtCore import QThread, pyqtSignal
from pathlib import Path
import docx
import os
from pdfminer.high_level import extract_text
from collections import defaultdict
import jieba
from difflib import get_close_matches
import multiprocessing
from functools import partial
import time
import mmap
import psutil
from cache_manager import CacheManager
import logger_config

global logger
logger=logger_config.setup_logger(__name__)

def process_document(file_path, doc_type, timeout=180, cache_manager=None):  # 减少默认超时时间
    try:
        logger.info(f"\n[文档处理] 开始处理{doc_type}文档")
        logger.info(f"[文档处理] 文件路径: {file_path}")
        start_time = time.time()

        # 检查缓存
        if cache_manager:
            cached_doc = cache_manager.get_cached_document(file_path)
            if cached_doc:
                logger.info(f"[文档处理] 找到缓存数据，直接使用缓存")
                return cached_doc
        text = ""

        # 检查文件大小并设置动态超时
        file_size = os.path.getsize(file_path) / (1024 * 1024)  # 转换为MB
        logger.info(f"[文档处理] 文件大小: {file_size:.2f}MB")
        if file_size > 100:  # 如果文件大于100MB
            logger.info(f"[文档处理] 警告: 文件大小超过100MB，可能需要较长处理时间")
            timeout = min(timeout * 2, 600)  # 大文件增加超时时间，但不超过10分钟

        # 获取当前内存使用情况
        process = psutil.Process()
        memory_info = process.memory_info()
        initial_memory = memory_info.rss / 1024 / 1024
        logger.info(f"[文档处理] 初始内存使用: {initial_memory:.2f}MB")

        # 添加内存监控函数
        def check_memory_usage():
            current_memory = process.memory_info().rss / 1024 / 1024
            memory_increase = current_memory - initial_memory
            if memory_increase > 1024:  # 如果内存增加超过1GB
                raise MemoryError(f"内存使用增加过多: {memory_increase:.2f}MB")
            if time.time() - start_time > timeout:
                raise TimeoutError(f"处理超时（{timeout}秒）")

        if doc_type == 'docx':
            logger.info("\n[Word文档处理] 开始读取文件...")
            logger.info(f"[Word文档处理] 文件路径: {file_path}")
            # 使用内存映射优化Word文档读取
            with open(file_path, 'rb') as f:
                logger.info("[Word文档处理] 正在加载文档对象...")
                doc = docx.Document(f)
                text_parts = []
                total_paragraphs = len(doc.paragraphs)
                logger.info(f"[Word文档处理] 文档加载完成，总段落数: {total_paragraphs}")
                
                for i, paragraph in enumerate(doc.paragraphs):
                    if time.time() - start_time > timeout:
                        raise TimeoutError(f"处理超时（{timeout}秒）")
                    if paragraph.text.strip():  # 只添加非空段落
                        text_parts.append(paragraph.text)
                    if i % 50 == 0:  # 每处理50个段落输出一次进度
                        progress = (i / total_paragraphs) * 100
                        elapsed_time = time.time() - start_time
                        memory_info = process.memory_info()
                        logger.info(f"[Word文档处理] 进度: {progress:.1f}% ({i}/{total_paragraphs} 段落)")
                        logger.info(f"[Word文档处理] 已用时间: {elapsed_time:.1f}秒")
                        logger.info(f"[Word文档处理] 当前内存使用: {memory_info.rss / 1024 / 1024:.1f}MB")
                        if i > 0:
                            avg_time_per_para = elapsed_time / i
                            remaining_paras = total_paragraphs - i
                            estimated_remaining_time = avg_time_per_para * remaining_paras
                            logger.info(f"[Word文档处理] 预计剩余时间: {estimated_remaining_time:.1f}秒")
                
                text = '\n'.join(text_parts)
                logger.info(f"\n[Word文档处理] 文档读取完成")
                logger.info(f"[Word文档处理] 有效段落数: {len(text_parts)}")
                logger.info(f"[Word文档处理] 文本总长度: {len(text)} 字符")
                del doc, text_parts  # 释放内存
        else:  # pdf
            logger.info("\n[PDF文档处理] 开始读取文件...")
            logger.info(f"[PDF文档处理] 文件路径: {file_path}")
            # 使用pdfminer.six优化PDF处理
            # 检查文件大小，对大文件采取特殊处理
            if file_size > 50:  # 如果文件大于50MB
                logger.info(f"[PDF文档处理] 大文件处理模式")
                timeout = min(timeout, 300)  # 限制最大超时时间为5分钟

            logger.info("[PDF文档处理] 正在提取文本...")
            extract_start_time = time.time()
            try:
                with open(file_path, 'rb') as pdf_file:
                    text = extract_text(
                        pdf_file,
                        maxpages=50,  # 限制最大页数
                        caching=True,  # 启用缓存
                        codec='utf-8'
                    )
                    if not text:
                        raise ValueError("无法提取文本内容")
                    
                extract_time = time.time() - extract_start_time
                logger.info(f"[PDF文档处理] 文本提取耗时: {extract_time:.1f}秒")
                logger.info(f"[PDF文档处理] 文本提取完成，文本长度: {len(text)} 字符")
            except Exception as e:
                logger.info(f"[PDF文档处理] 文本提取失败: {str(e)}")
                return None

        logger.info("\n[分词处理] 开始进行分词...")
        text_lower = text.lower()
        logger.info(f"[分词处理] 文本预处理完成，准备分词")
        batch_size = min(5000, max(1000, int(1000000 / file_size)))  # 根据文件大小动态调整批处理大小
        words = []
        word_positions = defaultdict(list)
        
        # 使用生成器优化内存使用
        def word_generator(text):
            logger.info("[分词处理] 初始化分词生成器...")
            for i, word in enumerate(jieba.cut(text)):
                if i % 1000 == 0:  # 增加检查频率
                    check_memory_usage()
                yield i, word

        # 分批处理分词
        word_gen = word_generator(text_lower)
        current_batch = []
        total_words = 0
        batch_count = 0

        logger.info("[分词处理] 开始批量处理...")
        segment_start_time = time.time()
        while True:
            try:
                if time.time() - start_time > timeout:
                    raise TimeoutError(f"处理超时（{timeout}秒）")
                
                i, word = next(word_gen)
                current_batch.append((i, word))
                
                if len(current_batch) >= batch_size:
                    batch_count += 1
                    batch_start_time = time.time()
                    
                    for idx, w in current_batch:
                        words.append(w)
                        word_positions[w].append(idx)
                    total_words += len(current_batch)
                    
                    batch_time = time.time() - batch_start_time
                    segment_time = time.time() - segment_start_time
                    memory_info = process.memory_info()
                    
                    logger.info(f"\n[分词处理] 批次 {batch_count} 处理完成:")
                    logger.info(f"[分词处理] - 本批处理耗时: {batch_time:.2f}秒")
                    logger.info(f"[分词处理] - 总耗时: {segment_time:.2f}秒")
                    logger.info(f"[分词处理] - 当前已处理: {total_words} 个词")
                    logger.info(f"[分词处理] - 当前词典大小: {len(word_positions)} 个不同词")
                    logger.info(f"[分词处理] - 当前内存使用: {memory_info.rss / 1024 / 1024:.1f}MB")
                    logger.info(f"[分词处理] - 平均处理速度: {total_words/segment_time:.1f} 词/秒")
                    
                    current_batch = []
                    
            except StopIteration:
                # 处理最后一批
                if current_batch:
                    for idx, w in current_batch:
                        words.append(w)
                        word_positions[w].append(idx)
                    total_words += len(current_batch)
                    logger.info("\n[分词处理] 处理最后一批数据完成")
                break

        process_time = time.time() - start_time
        memory_info = process.memory_info()
        
        logger.info(f"\n[处理完成] 文档处理结果汇总:")
        logger.info(f"[处理完成] - 总用时: {process_time:.2f}秒")
        logger.info(f"[处理完成] - 总词数: {len(words)}")
        logger.info(f"[处理完成] - 不同词数: {len(word_positions)}")
        logger.info(f"[处理完成] - 最终内存使用: {memory_info.rss / 1024 / 1024:.1f}MB")
        logger.info(f"[处理完成] - 平均处理速度: {len(words)/process_time:.1f} 词/秒")

        result = {
            'path': str(file_path),
            'content': text,
            'type': doc_type,
            'word_positions': dict(word_positions),
            'words': words,
            'process_time': process_time
        }

        # 保存到缓存
        if cache_manager:
            cache_manager.cache_document(file_path, result)
            logger.info(f"[文档处理] 文档处理结果已保存到缓存")

        return result
    except Exception as e:
        logger.info(f"\n错误: 处理{doc_type}文档 {file_path} 失败")
        logger.info(f"错误信息: {str(e)}")
        logger.info(f"错误类型: {type(e).__name__}")
        return None

class DocumentScanner(QThread):
    progress_updated = pyqtSignal(int)
    scan_completed = pyqtSignal(tuple)

    def __init__(self, directory):
        super().__init__()
        self.directory = directory
        self.inverted_index = defaultdict(list)
        self.documents = []
        self.cpu_count = multiprocessing.cpu_count()
        self.cache_manager = CacheManager()

    def build_inverted_index(self, documents):
        inverted_index = defaultdict(list)
        for doc_id, doc in enumerate(documents):
            if doc is None:
                continue
            for pos, word in enumerate(doc['words']):
                inverted_index[word].append((doc_id, pos))
        return inverted_index

    def run(self):
        logger.info("\n开始扫描文档...")
        start_time = time.time()
        self.documents = []
        self.inverted_index.clear()
        
        # 获取所有文档路径
        docx_files = list(Path(self.directory).rglob('*.docx'))
        pdf_files = list(Path(self.directory).rglob('*.pdf'))
        total_files = len(docx_files) + len(pdf_files)
        
        logger.info(f"找到 {len(docx_files)} 个Word文档和 {len(pdf_files)} 个PDF文档")
        
        if total_files == 0:
            logger.info("未找到任何文档")
            self.scan_completed.emit(([], defaultdict(list)))
            return
        self.cpu_count = min(self.cpu_count, 4)  # 限制最大进程数
        batch_size = 10  # 每批处理的文件数
        
        # 分批处理Word文档
        docx_results = []
        logger.info("\n开始处理Word文档...")
        for i in range(0, len(docx_files), batch_size):
            batch_files = docx_files[i:i + batch_size]
            logger.info(f"处理Word文档批次 {i//batch_size + 1}/{(len(docx_files)-1)//batch_size + 1}")
            with multiprocessing.Pool(processes=self.cpu_count) as pool:
                batch_results = list(pool.imap(partial(process_document, doc_type='docx', cache_manager=self.cache_manager), batch_files))
                docx_results.extend([r for r in batch_results if r])
                progress = min(50, int((i + len(batch_files)) / total_files * 50))
                self.progress_updated.emit(progress)
        # 分批处理PDF文档
        pdf_results = []
        logger.info("\n开始处理PDF文档...")
        for i in range(0, len(pdf_files), batch_size):
            batch_files = pdf_files[i:i + batch_size]
            logger.info(f"处理PDF文档批次 {i//batch_size + 1}/{(len(pdf_files)-1)//batch_size + 1}")
            with multiprocessing.Pool(processes=self.cpu_count) as pool:
                batch_results = list(pool.imap(partial(process_document, doc_type='pdf', cache_manager=self.cache_manager), batch_files))
                pdf_results.extend([r for r in batch_results if r])
                progress = 50 + min(50, int((i + len(batch_files)) / total_files * 50))
                self.progress_updated.emit(progress)
        # 合并结果
        self.documents = docx_results + pdf_results
        self.inverted_index = self.build_inverted_index(self.documents)
        # 清理临时数据
        for doc in self.documents:
            doc.pop('words', None)
        total_time = time.time() - start_time
        logger.info("文档扫描完成")
        logger.info(f"总用时: {total_time:.2f}秒")
        logger.info(f"成功处理: {len(self.documents)}/{total_files} 个文档")
        logger.info(f"索引词数量: {len(self.inverted_index)}")
        self.scan_completed.emit((self.documents, self.inverted_index))
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Word文档全文检索系统")
        self.setMinimumSize(800, 600)
        self.documents = []
        self.inverted_index = defaultdict(list)  # 初始化inverted_index
        self.setup_ui()

    def setup_ui(self):
        # 创建主窗口部件和布局
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)

        # 创建顶部控制栏
        control_layout = QHBoxLayout()
        self.folder_path = QLineEdit()
        self.folder_path.setPlaceholderText("选择文档所在文件夹（支持Word和PDF）")
        browse_button = QPushButton("浏览")
        browse_button.clicked.connect(self.browse_folder)
        control_layout.addWidget(self.folder_path)
        control_layout.addWidget(browse_button)

        # 创建搜索栏
        search_layout = QHBoxLayout()
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("输入搜索关键词")
        self.search_input.returnPressed.connect(self.search_documents)
        search_button = QPushButton("搜索")
        search_button.clicked.connect(self.search_documents)
        search_layout.addWidget(self.search_input)
        search_layout.addWidget(search_button)

        # 创建进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)

        # 创建结果显示区域
        self.results_display = QTextEdit()
        self.results_display.setReadOnly(True)

        # 添加所有组件到主布局
        layout.addLayout(control_layout)
        layout.addLayout(search_layout)
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.results_display)

    def browse_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "选择文件夹")
        if folder:
            self.folder_path.setText(folder)
            self.scan_documents(folder)

    def scan_documents(self, folder):
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.scanner = DocumentScanner(folder)
        self.scanner.progress_updated.connect(self.update_progress)
        self.scanner.scan_completed.connect(self.scan_finished)
        self.scanner.start()

    def update_progress(self, value):
        self.progress_bar.setValue(value)

    def scan_finished(self, scan_result):
        self.documents, self.inverted_index = scan_result
        self.progress_bar.setVisible(False)
        docx_count = sum(1 for doc in self.documents if doc.get('type') == 'docx')
        pdf_count = sum(1 for doc in self.documents if doc.get('type') == 'pdf')
        self.results_display.setText(f"已扫描 {docx_count} 个Word文档和 {pdf_count} 个PDF文档")

    def search_documents(self):
        keyword = self.search_input.text().lower()
        if not keyword:
            return

        logger.info(f"开始搜索关键词: {keyword}")
        start_time = time.time()

        # 显示进度条
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)

        # 使用结巴分词处理搜索关键词
        keywords = jieba.lcut(keyword)
        logger.info(f"分词结果: {', '.join(keywords)}")
        results = []
        doc_scores = defaultdict(float)
        total_docs = len(self.documents)

        # 对每个分词后的关键词进行搜索
        for word in keywords:
            # 支持模糊匹配，但限制相似词数量
            similar_words = get_close_matches(word, self.inverted_index.keys(), n=2, cutoff=0.8)
            logger.info(f"处理关键词: {word}, 找到相似词: {', '.join(similar_words)}")
            for similar_word in similar_words:
                matches = self.inverted_index[similar_word]
                logger.info(f"相似词 '{similar_word}' 在 {len(set(doc_id for doc_id, _ in matches))} 个文档中找到匹配")
                for doc_id, pos in matches:
                    # 根据词频和位置计算文档得分
                    doc = self.documents[doc_id]
                    freq = len(doc['word_positions'][similar_word])
                    position_score = 1.0 / (pos + 1)  # 关键词出现位置越靠前，得分越高
                    similarity_score = 1.0 if word == similar_word else 0.7  # 完全匹配得分更高
                    doc_scores[doc_id] += freq * position_score * similarity_score

            # 更新进度条
            progress = int((keywords.index(word) + 1) / len(keywords) * 40)
            self.progress_bar.setValue(progress)

        # 根据得分对文档排序
        scored_docs = [(doc_id, score) for doc_id, score in doc_scores.items()]
        scored_docs.sort(key=lambda x: x[1], reverse=True)

        logger.info(f"搜索完成，找到 {len(scored_docs)} 个匹配文档")
        logger.info(f"搜索用时: {time.time() - start_time:.2f}秒")

        # 分页处理，每页显示10个结果
        page_size = 10
        total_results = len(scored_docs)
        current_page = 0

        # 处理第一页结果
        page_docs = scored_docs[current_page:current_page + page_size]
        for doc_id, score in page_docs:
            doc = self.documents[doc_id]
            results.append(f"文件: {doc['path']}\n得分: {score:.2f}\n")
            
            # 获取关键词上下文，限制上下文数量
            content = doc['content'].lower()
            contexts = []
            for word in keywords:
                similar_words = get_close_matches(word, self.inverted_index.keys(), n=2, cutoff=0.8)
                for similar_word in similar_words:
                    positions = doc['word_positions'].get(similar_word, [])
                    # 只获取前两个位置的上下文
                    for pos in positions[:2]:
                        start = max(0, content.rfind(' ', 0, pos) - 20)
                        end = min(len(content), content.find(' ', pos + len(similar_word)) + 20)
                        context = content[start:end].strip()
                        contexts.append(context)
            
            if contexts:
                results.append(f"上下文:\n" + "\n...".join(contexts[:2]) + "\n\n")

            # 更新进度条
            progress = 40 + int((page_docs.index((doc_id, score)) + 1) / len(page_docs) * 40)
            self.progress_bar.setValue(progress)

        if results:
            result_text = ''.join(results)
            if total_results > page_size:
                result_text += f"\n--- 显示 {page_size}/{total_results} 个结果 ---\n"
            self.results_display.setText(result_text)
        else:
            self.results_display.setText("未找到匹配的结果")

        # 高亮显示关键词和文件路径
        cursor = self.results_display.textCursor()
        format = QTextCharFormat()
        format.setBackground(QColor("yellow"))

        # 先高亮文件路径
        cursor.movePosition(QTextCursor.MoveOperation.Start)
        while True:
            cursor = self.results_display.document().find("文件: ", cursor)
            if cursor.isNull():
                break
            cursor.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.MoveAnchor, 4)
            cursor.movePosition(QTextCursor.MoveOperation.EndOfLine, QTextCursor.MoveMode.KeepAnchor)
            cursor.mergeCharFormat(format)

        # 高亮所有关键词及其相似词
        cursor.movePosition(QTextCursor.MoveOperation.Start)
        for word in keywords:
            similar_words = get_close_matches(word, self.inverted_index.keys(), n=2, cutoff=0.8)
            for similar_word in similar_words:
                cursor.movePosition(QTextCursor.MoveOperation.Start)
                while True:
                    cursor = self.results_display.document().find(similar_word, cursor, 
                                                                 QTextDocument.FindFlag.FindCaseSensitively)
                    if cursor.isNull():
                        break
                    cursor.mergeCharFormat(format)

        # 再高亮关键词
        cursor.movePosition(QTextCursor.MoveOperation.Start)
        while cursor.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.KeepAnchor, len(keyword)):
            if cursor.selectedText().lower() == keyword:
                cursor.mergeCharFormat(format)
            cursor.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.MoveAnchor, -len(keyword) + 1)

        # 完成后隐藏进度条
        self.progress_bar.setVisible(False)

def main():
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())

if __name__ == '__main__':
    main()